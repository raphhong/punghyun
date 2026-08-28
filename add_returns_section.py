# -*- coding: utf-8 -*-
"""출자제안서에 '출자자 수익 — 얼마 넣으면 얼마 버나' 시뮬 섹션 삽입 (idempotent)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\sol hong\Downloads\풍현_출자제안서_260810.docx'
KFONT = '맑은 고딕'
NAVY = RGBColor(0x14,0x24,0x3B)
MONEY = RGBColor(0x12,0xA0,0x6D)
GREY = RGBColor(0x55,0x55,0x55)
WHITE = RGBColor(0xFF,0xFF,0xFF)

def setfont(run, size=None, bold=None, color=None):
    run.font.name = KFONT
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    rf.set(qn('w:eastAsia'), KFONT)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color

def set_borders(t):
    tblPr = t._element.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement('w:'+edge)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4')
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'CCCCCC')
        borders.append(e)
    tblPr.append(borders)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),hexcolor)
    tcPr.append(shd)

def set_cell(cell, text, size=9, bold=False, color=None, align='left', fill=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = {'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER,'right':WD_ALIGN_PARAGRAPH.RIGHT}[align]
    for i, line in enumerate(text.split('\n')):
        if i>0:
            p = cell.add_paragraph(); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.15
            p.alignment = {'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER,'right':WD_ALIGN_PARAGRAPH.RIGHT}[align]
        r = p.add_run(line); setfont(r, size, bold, color)
    if fill: shade(cell, fill)

d = Document(SRC)

# idempotency guard
for p in d.paragraphs:
    if p.text.strip().startswith('출자자 수익 — 얼마 넣으면'):
        print('already present; abort'); raise SystemExit

# reference: explanation para inserted after 투자구조표
ref = None
for p in d.paragraphs:
    if p.text.strip().startswith('풍현의 자금 조달은 성격이 다른 세 층으로'):
        ref = p; break
if ref is None:
    ref = d.paragraphs[-1]
    print('WARN: fallback ref = last para')

# --- build elements (appended at end, then relocated) ---
def new_para():
    return d.add_paragraph()

def heading(text):
    p = new_para(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3)
    r = p.add_run(text); setfont(r, 11, True, NAVY); return p

def small(text):
    p = new_para(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.25
    r = p.add_run(text); setfont(r, 9, False, GREY); return p

def subh(text):
    p = new_para(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2)
    r = p.add_run(text); setfont(r, 10, True, NAVY); return p

def build_table(headers, rows, widths, aligns, hl):
    t = d.add_table(rows=1, cols=len(headers)); t.autofit=False; set_borders(t)
    for j,h in enumerate(headers):
        set_cell(t.rows[0].cells[j], h, 9, True, WHITE, 'center', '14243B')
    for row in rows:
        cells = t.add_row().cells
        for j,val in enumerate(row):
            set_cell(cells[j], val, 9.5, bold=(j in hl), color=(MONEY if j in hl else None), align=aligns[j])
    for j,w in enumerate(widths):
        for r_ in t.rows:
            r_.cells[j].width = Cm(w)
    return t

seq = []
seq.append(heading('출자자 수익 — 얼마 넣으면 얼마 버나 (예시)'))
seq.append(small('아래는 이해를 돕기 위한 예시이며 확정 수익이 아닙니다. 회전·재투자 여부, 실제 취급 규모, 대손에 따라 달라질 수 있습니다.'))

seq.append(subh('① 유동화대금 (회전형) — 유동화 1회전(3개월)당 7.5%'))
seq.append(build_table(
    ['투자금','1회전 수익(7.5%)','연 수익(단리·회전마다 수령)','연 수익(재투자 복리, 약 33.5%)'],
    [['1억원','750만원','3,000만원 (연 30%)','약 3,355만원'],
     ['3억원','2,250만원','9,000만원 (연 30%)','약 1억 665만원'],
     ['5억원','3,750만원','1억 5,000만원 (연 30%)','약 1억 6,775만원']],
    [2.6,3.6,4.6,4.6], ['center','right','right','right'], hl={2,3}))
seq.append(small('연 30%(단리)는 회전 수익을 매 회전 인출하는 경우, 약 33.5%(복리)는 원금+수익을 4회 재투자하는 경우의 연 환산치입니다. (1.075⁴−1≈33.5%)'))

seq.append(subh('② 전환사채(CB) — 총 3억원 모집 / 연리 30%'))
seq.append(build_table(
    ['투자금','연 쿠폰(30%)','3년 누적 이자','추가 업사이드'],
    [['1억원','3,000만원','9,000만원','지분 전환권'],
     ['3억원','9,000만원','2억 7,000만원','지분 전환권'],
     ['5억원','1억 5,000만원','4억 5,000만원','지분 전환권']],
    [2.6,3.8,3.8,4.4], ['center','right','right','center'], hl={1,2}))
seq.append(small('CB는 총 3억원 모집(사모)이며 위는 배정 규모별 예시입니다. 전환권 행사 시 이자에 더해 기업가치 상승분을 지분으로 취득할 수 있습니다. 조건은 협의로 확정합니다.'))
seq.append(small('※ 본 수익은 목표·예상치로 원금·수익을 보장하지 않으며, 전문/적격투자자 대상 사모를 전제로 합니다.'))

# relocate right after ref, preserving order
anchor = ref._element
for obj in seq:
    el = obj._element
    anchor.addnext(el)
    anchor = el
print('inserted', len(seq), 'elements after ref')

try:
    d.save(SRC); print('saved', SRC)
except PermissionError:
    alt = SRC.replace('.docx','_v2.docx'); d.save(alt); print('locked -> saved', alt)
print('DONE')

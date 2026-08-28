# -*- coding: utf-8 -*-
"""자금 출자자(유동화 펀딩 / CB)용 IR 문서 생성
출력: 풍현_유동화출자_IR_260810.docx
- 실물 자산담보부 3개월 회전 채권 투자
- '얼마 넣으면 얼마 번다' 수익 시뮬레이션 (1/3/5억)
- 유동화 펀딩: 분기 7.5% (연 약 30%) / CB: 연리 약 30% + 전환권
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\sol hong\Downloads'
KFONT = '맑은 고딕'
NAVY = RGBColor(0x14, 0x24, 0x3B)
GREEN = RGBColor(0x1F, 0x5C, 0x4D)
MONEY = RGBColor(0x12, 0xA0, 0x6D)
GREY = RGBColor(0x55, 0x55, 0x55)

def base_doc():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = KFONT
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), KFONT)
    sec = doc.sections[0]
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
    return doc

def _setfont(run, size=None, bold=None, color=None, name=KFONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    rf.set(qn('w:eastAsia'), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color

def cover_title(doc, text, size=22, color=NAVY, after=2, before=0, align='left'):
    p = doc.add_paragraph()
    p.alignment = {'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER}[align]
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text); _setfont(r, size, True, color)
    return p

def h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); _setfont(r, 13, True, GREEN)
    return p

def h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); _setfont(r, 11, True, NAVY)
    return p

def para(doc, text, size=10, bold=False, color=None, after=6, before=0, align='left'):
    p = doc.add_paragraph()
    p.alignment = {'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER}[align]
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text); _setfont(r, size, bold, color)
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run('· ' + text); _setfont(r, size)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell(cell, text, size=9, bold=False, color=None, align='left', fill=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = {'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER,'right':WD_ALIGN_PARAGRAPH.RIGHT}[align]
    for i, line in enumerate(text.split('\n')):
        if i > 0:
            p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.15
            p.alignment = {'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER,'right':WD_ALIGN_PARAGRAPH.RIGHT}[align]
        r = p.add_run(line); _setfont(r, size, bold, color)
    if fill: shade(cell, fill)

def make_table(doc, headers, rows, widths_cm, body_size=9, aligns=None, hl_cols=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.autofit = False
    for j, h in enumerate(headers):
        set_cell(t.rows[0].cells[j], h, 9, True, RGBColor(0xFF,0xFF,0xFF), 'center', '14243B')
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            al = aligns[j] if aligns else 'left'
            col = MONEY if (hl_cols and j in hl_cols) else None
            bd = bool(hl_cols and j in hl_cols)
            set_cell(cells[j], val, body_size, bold=bd, color=col, align=al)
    for j, w in enumerate(widths_cm):
        for row in t.rows:
            row.cells[j].width = Cm(w)
    return t

def callout(doc, title_text, body_text, fill='EAF2EE', accent=GREEN):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'; t.autofit = False
    c = t.rows[0].cells[0]
    c.text = ''
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title_text); _setfont(r, 10.5, True, accent)
    p2 = c.add_paragraph(); p2.paragraph_format.line_spacing = 1.3
    r2 = p2.add_run(body_text); _setfont(r2, 9.5, False, NAVY)
    shade(c, fill)
    c.width = Cm(17.0)
    return t

def note(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
    r = p.add_run(text); _setfont(r, 8.5, False, GREY)
    return p

def spacer(doc, pts=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts)
    return p

def build():
    doc = base_doc()

    # ===== COVER =====
    cover_title(doc, '투자 제안서 (자금 출자자용)', 22, NAVY, after=2)
    cover_title(doc, '주식회사 풍현', 15, GREEN, after=2)
    para(doc, '실물 자산담보부 3개월 회전 채권 투자 — 세일앤렌탈백(SRB) 유동화', size=11, bold=True, color=NAVY, after=2)
    para(doc, '유동화 펀딩(1회전 3개월당 7.5%) · 전환사채 CB(3억 모집, 연 30% 목표) — 실물 자산으로 뒷받침되는 수익 (예상치, 보장 아님)', size=10, color=GREY, after=2)
    para(doc, '2026 · CONFIDENTIAL · 전문/적격투자자 대상 · 무단 배포 금지', size=9, color=GREY, after=4)

    # ===== Ⅰ. Executive Summary =====
    h1(doc, 'Ⅰ. 한눈에 보는 제안 (Executive Summary)')
    para(doc, '주식회사 풍현은 「사업자가 가진 물건(장비)을 사서, 그 사업자에게 다시 빌려주는」 사업을 합니다. '
              '당장 현금이 필요한 사업자가 자기 장비를 풍현에 팔아 즉시 목돈을 마련하고, 그 장비를 그대로 계속 쓰면서 3개월간 사용료(렌탈료)를 냅니다. '
              '3개월 뒤 사업자는 장비를 되사가거나, 필요 없으면 반납할 수 있습니다. '
              '이렇게 하면 사업자는 당장 필요한 현금을 얻고, 풍현은 「실물 장비」를 손에 쥔 채 사용료 수익을 얻습니다. '
              '이 구조를 세일앤렌탈백(SRB, ‘팔고 다시 빌려쓰기’)이라고 부릅니다.',
         size=9.5, after=5)
    para(doc, '출자자(투자자)는 이 거래에 들어가는 자금을 대는 사람입니다. 풍현이 장비를 사려면 목돈이 필요한데, 그 자금을 출자자가 대고 3개월마다 수익을 돌려받습니다. '
              '모든 거래 뒤에는 실제로 되팔 수 있는 「실물 장비」가 있고 그 소유권을 풍현이 가지고 있어, 이것이 원금을 지키는 안전판이 됩니다.',
         size=9.5, after=6)

    callout(doc, '예를 들면 (아주 쉽게)',
            '식당 사장님이 3,000만 원짜리 주방설비를 갖고 있는데 당장 현금이 급합니다. 풍현이 이 설비를 제값(시세)에 사 줍니다. '
            '사장님은 설비를 그대로 쓰면서 3개월간 사용료를 냅니다. 3개월 뒤 사장님은 설비를 되사가거나(대부분 되사감), 필요 없으면 반납합니다. '
            '→ 사장님은 현금을 얻고, 풍현은 되팔 수 있는 설비를 손에 쥔 채 사용료를 벌고, 출자자는 그 설비를 사는 데 쓴 자금을 대고 수익을 받습니다.',
            fill='FBF4E6', accent=NAVY)
    spacer(doc, 4)

    make_table(doc,
        ['핵심 포인트', '내용'],
        [['실물 자산담보', '모든 채권 뒤에 실물 자산이 존재하고, 풍현이 시가에 매입해 소유권까지 보유'],
         ['단기 회전 (3개월)', '자금이 오래 묶이지 않음. 분기마다 원금 회수 → 재투자로 복리 효과'],
         ['목표 수익', '유동화 펀딩: 1회전(3개월)당 7.5% (연 환산 약 30%)  ·  CB: 3억 모집, 연리 약 30% + 전환권'],
         ['검증된 운영', '대표는 BLQ에서 월 최대 40억원 규모의 채권·정산 흐름을 실제 운영'],
         ['다층 안전구조', '자산 소유권 보유 + 정산 통제 + 비소구 설계로 원금의 하방을 방어']],
        [4.0, 13.0], body_size=9.5)

    # ===== Ⅱ. 안전성 =====
    h1(doc, 'Ⅱ. 왜 안전한가 — 원금을 지키는 3중 구조')
    para(doc, '자산 금융에서 가장 중요한 것은 수익률이 아니라 ‘떼이지 않는 것’입니다. 출자자의 자금은 아래 3중 구조로 보호됩니다.', size=9.5, after=4)
    h2(doc, '① 실물 자산 담보 + 소유권 보유')
    para(doc, '모든 채권 뒤에는 실물 자산이 존재합니다. 풍현은 이 자산을 시가에 상응하는 공정가격으로 매입하여 소유권까지 보유하므로, 고객이 미납하더라도 자산 자체가 회수원이 됩니다.', size=9.5, after=4)
    h2(doc, '② 단기 회전 (3개월 만기)')
    para(doc, '각 거래는 3개월 완결 구조입니다. 자금이 장기간 묶이지 않아 시장 변화에 대한 노출이 최소화되고, 분기마다 원금이 회수되어 재투자·회수 판단이 용이합니다.', size=9.5, after=4)
    h2(doc, '③ 정산 통제 + 비소구 설계')
    para(doc, '카드매출 등 정산 흐름을 통제하고, 미납 시 자산 회수·매각으로 원금의 대부분을 회수하도록 설계되어 있습니다. 대표가 BLQ에서 검증한 회수 엔진이 그대로 적용됩니다.', size=9.5, after=6)

    h2(doc, '대손 가정 (보수적)')
    make_table(doc,
        ['단계', '가정', '근거'],
        [['① 연체율(gross)', '취급액의 약 5%', '자산 보유·상환력 확인 사업자 선별심사 반영'],
         ['② 담보 회수율', '원금의 약 90%', '시가 매입 자산의 소유권 보유 → 재매각·인수로 잔존가치 회수'],
         ['③ 건당 순손실', '원금의 약 10%', '중고 처분가 변동·처분 비용·기간을 보수적 반영'],
         ['포트폴리오 순대손율', '취급액의 약 0.5%', '① × ③ = 5% × 10%']],
        [4.2, 4.0, 8.8], body_size=9)
    callout(doc, '스트레스 시나리오에서도 흑자',
            '순대손율 약 0.5%는 회전 수익으로 충분히 흡수 가능합니다. 연체율이 2배(10%)로 오르고 중고 처분 실현가가 30% 추가 하락하는 스트레스 상황에서도 순대손율은 약 2% 수준에 머물러, ‘실물 자산을 소유권째 확보하고 만기 잔존가치로 회수’하는 구조가 손실의 하방을 구조적으로 막아 줍니다.')

    # ===== Ⅲ. 자금 흐름 =====
    h1(doc, 'Ⅲ. 당신의 자금이 어떻게 굴러가는가')
    para(doc, '출자자의 자금은 아래 3개월 사이클을 따라 순환하며, 회전마다 수익을 발생시킵니다.', size=9.5, after=4)
    make_table(doc,
        ['단계', '흐름 (쉽게)'],
        [['① 자금 투입', '출자자가 자금을 넣습니다 (유동화 자금 또는 CB)'],
         ['② 장비 매입', '풍현이 그 자금으로 사업자의 장비를 제값에 사고, 소유권을 가집니다'],
         ['③ 렌탈 시작', '장비를 사업자에게 3개월 빌려줍니다 → 받을 사용료(렌탈료)가 생깁니다'],
         ['④ 수익 지급', '이 「받을 사용료」를 자금화하여 출자자에게 수익(차익)을 지급합니다'],
         ['⑤ 만기 정산', '3개월 뒤 사업자가 장비를 되사가거나 반납합니다 → 장비·현금 회수'],
         ['⑥ 원금 회수·재투자', '넣었던 원금을 돌려받고, 다시 넣으면 다음 3개월 사이클 시작 (복리)']],
        [3.5, 13.5], body_size=9.5)

    # ===== Ⅳ. 투자 상품 =====
    h1(doc, 'Ⅳ. 투자 상품 — 두 가지 방식')
    make_table(doc,
        ['구분', '유동화 펀딩 (회전형)', '전환사채 (CB)'],
        [['성격', '렌탈채권 매입·유동화 자금', '풍현이 발행하는 전환사채'],
         ['모집 구조', '유동화 1회전(3개월)마다 사용하는\n단기자금 (수시·1회성)', '총 3억원 목표 (사모)'],
         ['수익', '1회전(3개월)당 7.5% 수수료\n(재투자 시 연 환산 약 30%)', '연리 약 30% 쿠폰'],
         ['업사이드', '회전 재투자에 따른 복리', '+ 지분 전환권 (기업가치 상승 참여)'],
         ['기간', '3개월 만기 (회전마다 원금 회수)', '만기 협의 (예: 1~2년)'],
         ['담보·배경', '실물 자산 배경의 유동화 채권', '회사 신용 + 전환권'],
         ['원금 회수', '3개월마다 원금 회수', '만기 상환 또는 지분 전환'],
         ['적합 투자자', '단기·안정 회전 수익 선호', '중기·업사이드 동반 선호']],
        [2.6, 7.2, 7.2], body_size=9)
    note(doc, '유동화 펀딩은 유동화가 일어날 때마다 3개월 단위로 사용하는 1회성 단기자금이며(회전 재투자는 선택), CB는 총 3억원 규모의 사모 발행입니다. 두 방식은 병행 가능하고, 규모·조건은 협의로 확정합니다.')

    # ===== Ⅴ. 수익 시뮬레이션 (핵심) =====
    h1(doc, 'Ⅴ. 수익 시뮬레이션 — 얼마 넣으면 얼마 버나')
    para(doc, '아래는 이해를 돕기 위한 예시이며, 확정 수익이 아닙니다. 회전·재투자 여부, 실제 취급 규모, 대손에 따라 달라질 수 있습니다.', size=9, color=GREY, after=4)

    h2(doc, '① 유동화 펀딩 (회전형) — 분기 7.5% 기준')
    make_table(doc,
        ['투자금', '분기 수익 (7.5%)', '연 수익 (단리·분기 수령)', '연 수익 (재투자 복리, 약 33.5%)'],
        [['1억원', '750만원', '3,000만원 (연 30%)', '약 3,355만원'],
         ['3억원', '2,250만원', '9,000만원 (연 30%)', '약 1억 665만원'],
         ['5억원', '3,750만원', '1억 5,000만원 (연 30%)', '약 1억 6,775만원']],
        [3.0, 4.0, 5.0, 5.0], body_size=9.5,
        aligns=['center','right','right','right'], hl_cols=[2,3])
    note(doc, '연 30%(단리)는 분기 수익을 매 분기 인출하는 경우, 약 33.5%(복리)는 원금+수익을 4회 재투자하는 경우의 연 환산치입니다. (1.075⁴ − 1 ≈ 33.5%)')

    h2(doc, '② 전환사채(CB) — 연리 30% 기준')
    make_table(doc,
        ['투자금', '연 쿠폰 (30%)', '3년 누적 이자', '추가 업사이드'],
        [['1억원', '3,000만원', '9,000만원', '지분 전환권'],
         ['3억원', '9,000만원', '2억 7,000만원', '지분 전환권'],
         ['5억원', '1억 5,000만원', '4억 5,000만원', '지분 전환권']],
        [3.0, 4.5, 4.5, 5.0], body_size=9.5,
        aligns=['center','right','right','center'], hl_cols=[1,2])
    note(doc, 'CB는 총 3억원 모집(사모)을 목표로 하며, 위는 배정 규모별 예시입니다. 전환권 행사 시 이자 수익에 더해 기업가치 상승분을 지분으로 취득할 수 있습니다. 쿠폰율·전환조건·만기는 협의로 확정합니다.')

    callout(doc, '핵심 요약',
            '유동화 펀딩은 실물 자산으로 뒷받침되는 3개월 회전 채권에 자금을 대고 분기 7.5%(연 약 30%)를 수취하는 구조이며, CB는 연리 약 30% 고정 수익에 지분 전환 업사이드를 더한 구조입니다. '
            '두 방식 모두 풍현이 소유권을 보유한 실물 자산과 검증된 회수 엔진이 원금의 하방을 방어합니다.')

    # ===== Ⅵ. 대표/회사 신뢰 =====
    h1(doc, 'Ⅵ. 왜 풍현인가 — 검증된 운영 주체')
    bullet(doc, '엑시트 경험: 대표는 주식회사 비엘큐(BLQ)를 창업·운영하여 자비스앤빌런즈에 매각, 인수 후에도 대표이사로 사업을 이어 운영.')
    bullet(doc, '동종 실전 운영: BLQ에서 렌탈·정산·채권 회수를 직접 운영, 월 최대 40억원 규모의 채권·정산 흐름 관리 이력.')
    bullet(doc, '자금·구조 설계 역량: 채권양수도·유동화·포괄약정·정산 자동화 등 본 사업에 필요한 금융 구조를 실무에서 설계·운영.')
    bullet(doc, '전략적 투자 유치: 임팩트 투자사 소풍벤처스가 지분 투자로 참여하여 초기부터 신뢰 기반 확보.')

    # ===== Ⅶ. 리스크 및 대응 =====
    h1(doc, 'Ⅶ. 리스크 및 대응')
    make_table(doc,
        ['리스크', '대응'],
        [['고객 미납', '자산 회수·매각(소유권 보유) + 카드매출 등 정산 통제'],
         ['중고 처분가 하락', '시가 상응 매입으로 확보한 매입가 버퍼, 보수적 잔존가치(시가의 약 90%) 가정'],
         ['회전 지연', '3개월 단기 만기로 노출 최소화, 재투자 시점·규모 유연 조정'],
         ['규제(대부업·유사수신)', '진성매매+임대차 구조, 법률의견서 확보, 전문/적격투자자 대상 사모로 진행'],
         ['운영 리스크', '대표의 동종 사업 실전 운영 역량 + 소풍벤처스 투자·거버넌스']],
        [4.0, 13.0], body_size=9.5)

    # ===== Ⅷ. 투자 실행 =====
    h1(doc, 'Ⅷ. 투자 실행')
    make_table(doc,
        ['항목', '내용'],
        [['최소 투자금', '1억원부터 (예시, 협의 가능)'],
         ['투자 방식', '유동화 펀딩(3개월 회전, 수시·1회성) / 전환사채(CB, 총 3억 모집) / 병행'],
         ['수익 지급', '유동화: 회전(3개월)별 정산  ·  CB: 약정 쿠폰 지급'],
         ['진행 절차', '① 협의·실사 → ② 계약 체결 → ③ 자금 투입 → ④ 회전·정산 → ⑤ 원금 회수/재투자'],
         ['대상 투자자', '전문투자자·적격투자자 (사모)']],
        [4.0, 13.0], body_size=9.5)

    para(doc, '주식회사 풍현  |  대표이사 홍솔', size=10, bold=True, color=NAVY, before=10, after=2)
    para(doc, '연락처:', size=9.5, color=GREY, after=6)

    note(doc, '본 문서는 전문/적격투자자를 위한 투자 검토용 참고 자료로, 불특정 다수에 대한 투자 권유나 수익 보장 문서가 아닙니다. 본 자료의 수익률(분기 7.5%·연 약 30% 등)은 예시 가정에 기반한 목표·예상치이며 확정 실적이나 원금을 보장하지 않습니다. '
                 '실제 투자 구조·수익·조건은 협의 및 계약으로 확정되며, 유사수신행위규제·자본시장법·대부업법 등 관련 법령 준수를 위해 실행 전 법률·세무·회계 검토를 거칩니다. 본 자료는 기밀이며 무단 배포를 금합니다.')

    path = OUT + r'\풍현_유동화출자_IR_260810.docx'
    doc.save(path); print('saved', path)

if __name__ == '__main__':
    build()
    print('DONE')

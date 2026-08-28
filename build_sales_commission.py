# -*- coding: utf-8 -*-
"""영업자용 수수료 제안서. 마진스택 비공개 — 기본급 2.07% + 표준(22.5%) 초과분 전액."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTDIR = r"C:\Users\sol hong\Downloads"
FONT = "맑은 고딕"
NAVY = RGBColor(0x0a, 0x18, 0x30)
GREEN = RGBColor(0x12, 0xa0, 0x6d)
GRAY = RGBColor(0x55, 0x5b, 0x66)
HDR_FILL = "0A1830"
BOX_FILL = "F2F6F4"
GREEN_FILL = "E7F5EE"
ALT_FILL = "F7F9FB"

doc = Document()
st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(10.5)
st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
sec = doc.sections[0]
sec.left_margin = Cm(2.3); sec.right_margin = Cm(2.3)
sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)

def _font(run, size=None, bold=False, color=None, italic=False):
    run.font.name = FONT; run.bold = bold; run.italic = italic
    if size: run.font.size = Pt(size)
    if color is not None: run.font.color.rgb = color
    rF = run._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:eastAsia", "w:ascii", "w:hAnsi"): rF.set(qn(a), FONT)

def para(text="", size=10.5, bold=False, color=None, align=None, sa=6, sb=0, italic=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(sa); p.paragraph_format.space_before = Pt(sb)
    if text: _font(p.add_run(text), size, bold, color, italic)
    return p

def runs(parts, size=10.5, align=None, sa=6, sb=0):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(sa); p.paragraph_format.space_before = Pt(sb)
    for t, b, *c in parts: _font(p.add_run(t), size, b, c[0] if c else None)
    return p

def heading(text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(5)
    _font(p.add_run(text), size, True, NAVY)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6"); b.set(qn("w:space"), "4"); b.set(qn("w:color"), "12A06D")
    pbdr.append(b); pPr.append(pbdr)
    return p

def bullets(items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
        if isinstance(it, str): _font(p.add_run(it), size)
        else:
            for t, b, *c in it: _font(p.add_run(t), size, b, c[0] if c else None)

def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _cell(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""; p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    if align is not None: p.alignment = align
    for i, ln in enumerate(str(text).split("\n")):
        if i > 0:
            p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(2)
            if align is not None: p.alignment = align
        _font(p.add_run(ln), size, bold, color)

def table(headers, rows, widths=None, size=10, aligns=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, bold=True, color=RGBColor(0xff,0xff,0xff), size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(t.rows[0].cells[i], HDR_FILL)
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            al = aligns[i] if aligns else None
            _cell(cells[i], val, bold=(i==0), size=size, align=al)
            if r_i % 2 == 1: _shade(cells[i], ALT_FILL)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def callout(title, lines, fill=BOX_FILL, title_color=NAVY):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]; _shade(c, fill); c.text = ""
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(4)
    _font(p.add_run(title), 11, True, title_color)
    for ln in lines:
        pp = c.add_paragraph(); pp.paragraph_format.space_after = Pt(3)
        if isinstance(ln, str): _font(pp.add_run(ln), 10)
        else:
            for tx, b, *col in ln: _font(pp.add_run(tx), 10, b, col[0] if col else None)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

C = WD_ALIGN_PARAGRAPH.CENTER
R = WD_ALIGN_PARAGRAPH.RIGHT

# ===== TITLE =====
para("풍현 영업 파트너 수수료 안내", size=22, bold=True, color=NAVY, align=C, sa=2)
para("가져온 만큼, 잘 받아온 만큼 — 투명하게 나눕니다", size=11.5, italic=True, color=GRAY, align=C, sa=10)

# ===== 한 줄 요약 =====
callout("한 줄 요약", [
 [("취급액의 2.07%가 기본. 여기에, 표준 렌탈료율(22.5%)을 넘겨 받아오면 그 초과분은 전액 당신 몫입니다.", True)],
], fill=GREEN_FILL, title_color=GREEN)

# ===== 수수료 구조 =====
heading("수수료는 두 가지로 구성됩니다")
para("① 기본 수수료 — 성사되는 모든 건", size=11, bold=True, color=NAVY, sb=4, sa=2)
runs([("취급액(자산 매입가) × ", False), ("2.07%", True, GREEN),
      ("  ·  건이 성사되면 무조건 발생하는 기본급입니다.", False)])
para("② 초과 성과급 — 표준보다 잘 받아왔을 때", size=11, bold=True, color=NAVY, sb=6, sa=2)
runs([("(이번 딜 렌탈료율 − 표준 22.5%) × 취급액 × ", False), ("100%", True, GREEN)])
bullets([
 [("표준 렌탈료율은 ", False), ("22.5%(3개월)", True), (". 이 기준으로 받아오면 기본 수수료만.", False)],
 [("표준을 넘겨 받아온 만큼(예: 23.5%, 24%, 25%)의 ", False), ("초과분은 전액 당신에게", True, GREEN), (" 지급됩니다.", False)],
 [("렌탈료 상한은 ", False), ("25%(3개월)", True), (". 성과급 계산은 이 상한까지 인정됩니다.", False)],
])

# ===== 예시 표 =====
heading("얼마를 버나 — 예시 (단위: 만원)")
para("취급액과 받아온 렌탈료율에 따라 아래처럼 계산됩니다.", size=10, color=GRAY, sa=4)
table(
 ["취급액", "기본 수수료\n(2.07%)", "초과 성과급\n@23.5%", "초과 성과급\n@25%(상한)", "총 수령\n(@25% 기준)"],
 [
  ["5억", "1,035", "500", "1,250", "2,285"],
  ["10억", "2,070", "1,000", "2,500", "4,570"],
  ["15억", "3,105", "1,500", "3,750", "6,855"],
  ["20억", "4,140", "2,000", "5,000", "9,140"],
 ],
 widths=[2.6, 3.0, 3.0, 3.2, 3.2],
 aligns=[C, R, R, R, R],
)
para("※ '총 수령'은 렌탈료를 상한(25%)까지 받아온 경우입니다. 표준(22.5%)으로 받아오면 기본 수수료만 발생합니다.", size=9, color=GRAY)

# ===== 지급 방식 =====
heading("언제 받나 — 지급 방식")
table(["구분", "지급 시점"], [
 ["기본 수수료", "성사 시 50% 선지급 + 만기 정산 시 나머지 50%"],
 ["초과 성과급", "만기에 실제 실현(정상 종료·회수) 확인 후 지급"],
], widths=[4.0, 12.0])
para("초과 성과급을 만기 실현 후에 지급하는 이유는, 무리하게 받아온 뒤 부실이 나면 회사와 파트너 모두 손해이기 때문입니다. 끝까지 건전하게 마무리된 건에 대해 확실히 드립니다.", size=10, color=GRAY)

# ===== 왜 이렇게 =====
heading("왜 이렇게 설계했나")
bullets([
 [("취급액만 늘리는 경쟁이 아니라, ", False), ("잘 받아온 성과", True, GREEN), ("에 보상합니다. 표준을 넘긴 만큼은 온전히 파트너 몫입니다.", False)],
 "기본급은 안정적으로, 성과급은 실력에 따라 무제한(상한 내)으로 — 두 축으로 동기부여합니다.",
 "부실 없이 끝까지 마무리한 건에 성과급을 지급하므로, 파트너와 회사의 이해가 완전히 일치합니다.",
])

# ===== 면책 =====
para("", sa=4)
para("※ 본 안내는 참고용이며, 실제 수수료율·조건·지급 방식은 개별 협의 및 계약으로 확정됩니다. 렌탈료율·상한 등 정책은 시장 상황과 컴플라이언스 검토에 따라 조정될 수 있습니다. 대외비 · 무단 배포 금지.", size=8.5, color=GRAY)

fn = "풍현_영업수수료_제안_260818.docx"
p = os.path.join(OUTDIR, fn)
try:
    doc.save(p); print("saved:", p)
except PermissionError:
    p = os.path.join(OUTDIR, "풍현_영업수수료_제안_260818_v2.docx")
    doc.save(p); print("locked -> saved:", p)

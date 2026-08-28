# -*- coding: utf-8 -*-
"""제안서 삽입용: 실제 적용 사례 — MRI 검사장비 세일앤렌탈백 (2.5억원)
풍현 투자 제안서 스타일(맑은 고딕·네이비 헤더·음영 표)에 맞춘 워드 조각.
출력: Downloads\풍현_MRI사례_260813_v2.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\sol hong\Downloads\풍현_제안서_보강자료_합본_260813.docx'
KFONT = '맑은 고딕'
NAVY = RGBColor(0x0A, 0x18, 0x30)
GREEN = RGBColor(0x12, 0xA0, 0x6D)
GREY = RGBColor(0x5A, 0x64, 0x72)
NAVY_HEX = '0A1830'
LIGHT_HEX = 'EAF3EF'   # 옅은 그린(강조 박스)


def base_doc():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = KFONT
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), KFONT)
    sec = doc.sections[0]
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    return doc


def _setfont(run, size=None, bold=None, color=None, name=KFONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    rf.set(qn('w:eastAsia'), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color


def h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); _setfont(r, 15, True, NAVY)
    return p


def h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); _setfont(r, 11.5, True, GREEN)
    return p


def para(doc, text, size=10, bold=False, color=None, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text); _setfont(r, size, bold, color)
    return p


def note(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text); _setfont(r, 8.5, False, GREY)
    return p


def bullet(doc, lead, rest='', size=10):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    r0 = p.add_run('· ' + lead); _setfont(r0, size, True, NAVY)
    if rest:
        r1 = p.add_run('  ' + rest); _setfont(r1, size, False)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def set_cell(cell, text, size=9.5, bold=False, color=None, align='left', fill=None, valign='center'):
    cell.text = ''
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign'); va.set(qn('w:val'), valign); tcPr.append(va)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
    for i, line in enumerate(text.split('\n')):
        if i > 0:
            p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.2
            p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
        r = p.add_run(line); _setfont(r, size, bold, color)
    if fill: shade(cell, fill)


def make_table(doc, headers, rows, widths_cm, header_align='center', body_aligns=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    t.autofit = False
    for j, h in enumerate(headers):
        set_cell(t.rows[0].cells[j], h, size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                 align=header_align, fill=NAVY_HEX)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            al = body_aligns[j] if body_aligns else ('left' if j == 0 else 'left')
            set_cell(cells[j], val, size=9.5, align=al)
    for j, w in enumerate(widths_cm):
        for row in t.rows:
            row.cells[j].width = Cm(w)
    return t


def callout(doc, title_text, body_text):
    """옅은 그린 강조 박스(1x1 표)."""
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    cell.width = Cm(17)
    shade(cell, LIGHT_HEX)
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(3); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(title_text); _setfont(r, 10.5, True, NAVY)
    p2 = cell.add_paragraph(); p2.paragraph_format.line_spacing = 1.4; p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body_text); _setfont(r2, 9.5, False, RGBColor(0x22, 0x2A, 0x35))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============================================================
# 본문
# ============================================================
def build():
    doc = base_doc()

    h1(doc, 'Ⅳ-1. 실제 적용 사례 — MRI 검사장비 세일앤렌탈백 (2.5억원)')
    para(doc,
         '아래는 풍현이 실제로 매입·집행하는 MRI 검사장비 건을 세일앤렌탈백 구조에 대입한 사례입니다. '
         '추상적 설명이 아니라, 동 구조가 현장에서 어떻게 작동하고 어떻게 안전하게 회수되는지를 구체적인 숫자로 보여드립니다.',
         size=10, after=8)

    # 사례 개요
    h2(doc, '사례 개요')
    make_table(
        doc,
        ['항목', '내용'],
        [
            ['대상 자산', '의료영상 진단장비(MRI 시스템) 1대'],
            ['대상 사업자', '영상의학과 의원 (건강보험 수가·카드매출이 매일 발생)'],
            ['자금 니즈', '단기 운영자금 약 2.5억원 (은행 한도 소진, 집행까지 며칠 내 필요)'],
            ['매입가(공정가격)', '2.5억원 — 동일 기종 중고 거래사례·감정가 범위 내에서 신속·현금 매입을 반영한 보수적 공정가'],
            ['렌탈 기간', '3개월'],
            ['만기 선택권', '반납형 또는 인수형 (계약 시점에 사업자가 선택)'],
        ],
        widths_cm=[5.5, 11.5],
        body_aligns=['left', 'left'],
    )

    # 거래 흐름
    h2(doc, '거래 흐름 (실제 숫자)')
    make_table(
        doc,
        ['단계', '이 사례에서 일어나는 일'],
        [
            ['① 자산 실사·가격 산정',
             '동일 기종 중고 시세와 감정가를 근거로 공정가격 2.5억원을 산정합니다. 신속 매입·현금 지급 조건을 반영해 시세 범위 내에서 보수적으로 책정하여, 회수 시 하방 버퍼를 확보합니다.'],
            ['② 매매·소유권 이전',
             'MRI의 소유권이 풍현으로 이전되고, 풍현은 사업자에게 2.5억원을 수일 내 지급합니다. 사업자는 이 시점에 목돈을 확보합니다.'],
            ['③ 렌탈(임대차) 개시',
             '같은 MRI를 사업자에게 3개월간 대여합니다. 환자 검사와 진료는 중단 없이 그대로 이어집니다. 바뀌는 것은 소유 명의뿐입니다.'],
            ['④ 이용료 수취',
             '3개월 총 이용료 약 2,000만원(매입가의 약 8%)을 수취합니다. 연 환산 수치는 커 보이지만, 사업자가 3개월간 실제 부담하는 총액은 원금 대비 한 자릿수 구간입니다.'],
            ['⑤ 만기 선택',
             '3개월 만기에 사업자는 반납 또는 인수 중 하나를 선택합니다(아래 표).'],
            ['⑥ 정산·회전 완결',
             '반납 시 재매각, 인수 시 인수대금 수취로 1회전이 완결되고, 회수한 자금은 다음 자산 매입에 재투입됩니다.'],
        ],
        widths_cm=[4.0, 13.0],
        body_aligns=['left', 'left'],
    )

    # 만기 시나리오
    h2(doc, '만기 시나리오 — 사업자의 선택')
    make_table(
        doc,
        ['구분', '반납형', '인수형'],
        [
            ['사업자의 선택', 'MRI를 풍현에 반납\n(추가 부담·위약금 없음)', '사전 확정 인수가로 재매입'],
            ['인수가', '해당 없음', '2.5억원\n(매입가와 동일한 정액, 계약 시 확정)'],
            ['사업자 순부담(3개월)', '이용료 약 2,000만원\n(MRI를 최종 처분한 것과 동일)', '수령 2.5억 − 이용료·인수가 지급\n= 이용료 약 2,000만원 (자산은 그대로 보유)'],
            ['풍현의 회수 경로', '반납받은 MRI를 중고시장에 매각\n(보수적 매입으로 매입가 이상 회수 여지) + 이용료', '인수대금 2.5억 + 이용료'],
        ],
        widths_cm=[3.6, 6.7, 6.7],
        body_aligns=['left', 'left', 'left'],
    )
    note(doc,
         '※ 인수가는 매입가에 요율·기간을 더해 산정한 값이 아니라, 만기 예상 잔존가치를 근거로 계약 시 확정한 정액입니다'
         '(조기 인수 감액·지연 증액 없음). MRI는 감가가 완만하여 3개월 잔존가치가 매입가 수준에 근접하므로, 인수가를 매입가와 동일한 정액으로 정합니다.')

    # 자산 선정 3원칙
    h2(doc, '왜 이 자산인가 — 풍현 자산 선정 3원칙과의 부합')
    make_table(
        doc,
        ['원칙', 'MRI 사례에서의 충족'],
        [
            ['자산확보 원칙\n(명확한 중고시세)',
             '의료영상장비는 기종별 중고 거래사례와 감정가가 형성되어 있어 공정가 산정이 가능합니다.'],
            ['유동성 원칙\n(빈번한 중고거래)',
             '의원·병원 개폐업과 장비 교체 수요로 중고 거래가 지속 발생하여, 반납 시 재매각 경로가 확보됩니다.'],
            ['현금흐름 원칙\n(명확한 카드매출)',
             '영상의학과는 건강보험 수가와 카드결제가 매일 정산되어, 이용료 수취와 정산 통제가 안정적입니다.'],
        ],
        widths_cm=[5.0, 12.0],
        body_aligns=['left', 'left'],
    )

    # 안전성 콜아웃
    h2(doc, '이 사례가 보여주는 안전성')
    callout(
        doc,
        '소유권을 넘겨받는 구조 그 자체에서 나오는 안전성',
        '동 건에서 사업자가 이용료를 미납하거나 폐업하더라도, MRI는 이미 풍현 소유이므로 담보권 실행이나 경매 없이 '
        '곧바로 회수·재매각할 수 있습니다. 매입가 수준에 근접하는 중고 잔존가치와 보수적으로 책정한 매입가가 원금의 하방을 이중으로 받쳐 주며, '
        '카드매출·수가 정산 흐름의 통제가 회수를 보조합니다(정산·수취 장치이며 담보가 아닙니다). '
        '회수는 자산과 정산 흐름에 한정되고 사업자의 다른 재산에 소구하지 않습니다. '
        '즉 이 거래의 안전성은 ‘좋은 고객’이 아니라 ‘소유권을 넘겨받는 구조 그 자체’에서 나옵니다.')

    # ========================================================
    # 추가 1: Ⅲ-③ 자금·구조 설계 역량
    # ========================================================
    doc.add_page_break()
    h1(doc, "Ⅲ-③ 자금·구조 설계 역량 — ‘회수’를 구조로 통제하는 노하우")
    para(doc,
         '대표는 채권양수도·유동화·정산 자동화 등 이 사업에 필요한 금융 구조를 실무에서 직접 설계·운영해 왔습니다. '
         '이 사업은 결국 「떼이지 않는 것」이 전부이고, 대표의 노하우는 문제가 생긴 뒤 쫓아가는 것이 아니라 '
         '애초에 떼이지 않도록 거래 구조를 미리 짜 두는 데 있습니다. 구체적으로는 이렇습니다.',
         size=10, after=6)
    bullet(doc, '돈이 들어오기 전에 먼저 받습니다.',
           '사업자의 카드매출이 사업자에게 가기 전에, 미리 정한 계좌로 이용료가 먼저 들어오도록 설계합니다. (담보가 아니라 정산 장치입니다.)')
    bullet(doc, '잘 팔리는 자산만 삽니다.',
           '중고로 빠르고 안정적으로 팔리는 물건인지 시세 데이터로 확인하고 매입합니다. 회수 안전은 살 때 이미 정해집니다.')
    bullet(doc, '자금을 돌리는 구조를 직접 짭니다.',
           '채권양수도·유동화·포괄약정 등 자금이 계속 회전하도록 만드는 구조를 실무로 설계·운영해 왔습니다.')
    bullet(doc, '카드매출로 미리 거릅니다.',
           '카드매출 흐름을 보고 실제로 갚을 여력이 있는 곳부터 골라 부실 가능성을 낮춥니다.')
    bullet(doc, '여러 건을 자동으로 관리합니다.',
           '정산을 자동화해, 취급 규모가 커져도 회수 관리가 흔들리지 않습니다.')
    para(doc,
         '모두 대표가 BLQ에서 월 최대 40억원 규모의 채권·정산을 실제로 운영하며 검증한 방식이며, 풍현은 이를 그대로 가져옵니다.',
         size=10, before=4, after=6)

    # ========================================================
    # 추가 2: 대상 사업자 — 업종 예시
    # ========================================================
    h1(doc, '대상 사업자 — 누가 대상인가')
    para(doc,
         '쉽게 말해, 팔면 값이 분명한 장비를 가지고 있고 카드매출이 꾸준한 사업자입니다. '
         '단기 자금은 필요한데 은행 문턱은 높은 곳들이죠. 예를 들면 다음과 같습니다.',
         size=10, after=6)
    make_table(
        doc,
        ['업종', '대표 장비·자산', '왜 적합한가'],
        [
            ['병·의원(영상의학과·정형외과 등)', 'MRI·CT·초음파 등', '장비가 비싸고 중고값이 분명 + 카드·수가 매일 발생'],
            ['치과', '유닛체어·구강스캐너·치과 CT', '고가 장비 + 카드결제 꾸준'],
            ['피부과·성형·에스테틱', '레이저·미용 의료기기', '장비 교체가 잦아 중고거래 활발 + 카드매출 큼'],
            ['동물병원', '진단·영상·수술 장비', '카드매출 명확 + 장비 중고시장 존재'],
            ['외식·카페·베이커리', '주방설비·오븐·냉장설비 등', '카드매출 명확 + 설비 중고시장이 두터움'],
            ['피트니스·필라테스·골프연습장', '운동기구·리포머·스크린골프 장비', '회원 카드결제 안정 + 장비 중고거래 활발'],
            ['물류·운송·건설', '차량·지게차·굴착기 등', '등록으로 소유권이 분명 + 중고시장이 큼'],
        ],
        widths_cm=[4.5, 5.5, 7.0],
        body_aligns=['left', 'left', 'left'],
    )
    para(doc,
         '공통점은 하나입니다 — 팔기 쉬운 자산 + 꾸준한 카드매출. 그래서 대상은 「업종」이 아니라 '
         '「자산이 잘 팔리는가, 카드매출이 꾸준한가」로 고릅니다.',
         size=10, before=6, after=4)

    doc.save(OUT)
    print('saved:', OUT)


if __name__ == '__main__':
    build()

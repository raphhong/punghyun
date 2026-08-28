# -*- coding: utf-8 -*-
"""풍현_출자제안서_260810.docx 에 3층 자금조달 구조 반영
① 지분 출자(소풍 2억/35%) ② CB 3억(연30%) ③ 단기 유동화자금(회전당 7.5%)
사용자가 수동 편집한 원본을 로드해 targeted edit 후 저장.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\sol hong\Downloads\풍현_출자제안서_260810.docx'
KFONT = '맑은 고딕'
NAVY = RGBColor(0x14, 0x24, 0x3B)
GREY = RGBColor(0x55, 0x55, 0x55)

def setfont(run, size=None, bold=None, color=None):
    run.font.name = KFONT
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    rf.set(qn('w:eastAsia'), KFONT)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color

def replace_text(p, new):
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new)

def set_cell(cell, new):
    p0 = cell.paragraphs[0]
    replace_text(p0, new)
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)

d = Document(SRC)

# ---- Edit 1: Executive Summary 요청사항 문단 ----
done1 = False
for p in d.paragraphs:
    if p.text.strip().startswith('본 제안은 풍현 설립을 위한'):
        replace_text(p,
            '본 제안의 자금 조달은 성격이 다른 세 층으로 구성됩니다. '
            '① 설립 지분 출자 — 소풍벤처스 2억원(지분 35%), '
            '② 성장 조달 — 전환사채(CB) 3억원(연 30% 목표, 지분 전환권), '
            '③ 회전 운영자금 — 유동화대금(유동화 1회전·3개월당 7.5% 수수료, 회전마다 사용하는 단기자금). '
            '상세는 Ⅸ장에 기술합니다.')
        done1 = True
        break
print('edit1 exec-summary:', done1)

# ---- Edit 2: 가정표(TABLE 8) 유동화 수수료율 행 주석 ----
done2 = False
t8 = d.tables[8]
for row in t8.rows:
    if row.cells[0].text.strip().startswith('유동화 수수료율'):
        set_cell(row.cells[0], '유동화 수수료율 (조달비용)')
        set_cell(row.cells[1], '채권액면의 7.5% / 3개월 — 단기 유동화자금 출자자 몫(연 환산 약 30%)')
        done2 = True
        break
print('edit2 assumption-row:', done2)

# ---- Edit 3: 투자 구조표(TABLE 12) 자금조달 3층으로 재구성 ----
t12 = d.tables[12]
rows = t12.rows
# rows[0]=헤더(항목|내용), rows[1..5]=데이터 5행
newvals = [
    ('① 설립 지분 출자 (소풍)', '2억원 / 지분 35% (보통주) · Post-money 약 5.7억원'),
    ('   대표 지분', '65%'),
    ('② 성장 조달 — 전환사채(CB)', '3억원 모집 / 연 30% 목표 / 지분 전환권 (만기 협의)'),
    ('③ 회전 운영자금 — 유동화대금', '유동화 1회전(3개월)당 7.5% 수수료 / 회전마다 사용하는 단기자금'),
    ('자금 성격', '지분(설립) + 부채성(CB·유동화). 구체 조건은 협의로 확정'),
]
for i, (a, b) in enumerate(newvals, start=1):
    set_cell(rows[i].cells[0], a)
    set_cell(rows[i].cells[1], b)
print('edit3 structure-table rows set:', len(newvals))

# ---- Edit 4: TABLE 12 뒤에 3층 설명 문단 삽입 ----
exp = d.add_paragraph()  # append at end, then move
exp.paragraph_format.space_before = Pt(4); exp.paragraph_format.space_after = Pt(6)
exp.paragraph_format.line_spacing = 1.3
r = exp.add_run(
    '풍현의 자금 조달은 성격이 다른 세 층으로 구성됩니다. ① 지분 출자(소풍 2억·35%)는 설립 자본이며 배당·지분가치로 회수합니다. '
    '② 전환사채(CB) 3억원은 성장 조달로, 연 30% 목표 수익과 지분 전환권을 제공합니다. '
    '③ 유동화대금은 유동화 1회전(3개월)마다 사용하는 회전 운영자금으로, 회전당 7.5%(연 환산 약 30%)의 수수료를 지급합니다. '
    '유동화대금의 조달비용(유동화 수수료)은 재무 전망(Ⅷ장)에 조달비용으로 이미 반영되어 있으며, CB 이자는 조달 규모 확정 시 반영합니다.')
setfont(r, 9.5, False, NAVY)
t12._element.addnext(exp._element)
print('edit4 explanation inserted')

# ---- Edit 5: 자금 사용 계획 ③ 문구 보강 (유동화 파트너십 → 조달 다변화) ----
done5 = False
for p in d.paragraphs:
    if p.text.strip().startswith('③ 유동화 파트너십'):
        replace_text(p, '③ 조달 다변화: CB·유동화대금 등 조달 채널을 다변화하여 회전 자금의 안정성을 확보합니다.')
        done5 = True
        break
print('edit5 fund-use:', done5)

# ---- 저장 (lock 시 _v2) ----
try:
    d.save(SRC); print('saved', SRC)
except PermissionError:
    alt = SRC.replace('.docx', '_v2.docx')
    d.save(alt); print('locked -> saved', alt)
print('DONE')
